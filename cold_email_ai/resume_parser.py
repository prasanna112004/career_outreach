"""
Resume upload parsing: extract plain text from PDF, DOCX, or TXT and
prepare LangChain documents for RAG.

For personal use, the user's resume becomes the knowledge base that is
matched semantically against the job description.
"""

from __future__ import annotations

import io
import re
from typing import List

from langchain.docstore.document import Document


def clean_resume_text(text: str) -> str:
    """Normalize whitespace and strip noise from extracted resume text."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(data: bytes) -> str:
    """Extract text from a PDF byte buffer."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: List[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            if t:
                parts.append(t)
        except Exception:
            continue
    return clean_resume_text("\n".join(parts))


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from a DOCX byte buffer."""
    import docx

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return clean_resume_text("\n".join(paragraphs))


def extract_text_from_txt(data: bytes) -> str:
    """Decode plain text; tries utf-8 then latin-1."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return clean_resume_text(data.decode(enc))
        except UnicodeDecodeError:
            continue
    return clean_resume_text(data.decode("utf-8", errors="replace"))


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """
    Route to the correct extractor based on file extension.

    Raises
    ------
    ValueError
        If the file type is not supported.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if lower.endswith(".docx"):
        return extract_text_from_docx(data)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return extract_text_from_txt(data)
    raise ValueError(
        "Unsupported file type. Please upload a PDF, DOCX, or TXT file."
    )


def guess_name_from_resume(text: str) -> str | None:
    """Heuristic: first substantial line is often the candidate name."""
    for line in text.split("\n"):
        line = line.strip()
        if 2 <= len(line) <= 80 and not re.match(
            r"^[\d\s\-\+\(\)@\.]+$", line
        ):
            # Skip lines that look like section headers only
            if line.isupper() and len(line) < 15:
                continue
            return line
    return None


def build_candidate_profile_summary(resume_text: str, max_chars: int = 1500) -> str:
    """
    Build a short profile block for the prompt (top of resume + optional
    summary section heuristics). Not a replacement for RAG chunks.
    """
    text = clean_resume_text(resume_text)
    if not text:
        return "No resume text extracted."

    name = guess_name_from_resume(text)
    head = text[:max_chars]
    parts = []
    if name:
        parts.append(f"Name (inferred from resume header): {name}")
    parts.append("Resume excerpt (for grounding):")
    parts.append(head)
    if len(text) > max_chars:
        parts.append("\n[... resume continues; see retrieved chunks below ...]")
    return "\n".join(parts)


def resume_text_to_documents(resume_text: str) -> List[Document]:
    """
    Wrap full resume text as source documents for chunking.

    Chunking is applied in the RAG pipeline (RecursiveCharacterTextSplitter).
    """
    text = clean_resume_text(resume_text)
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"section": "resume", "source": "user_upload"},
        )
    ]


__all__ = [
    "clean_resume_text",
    "extract_text_from_upload",
    "build_candidate_profile_summary",
    "resume_text_to_documents",
    "guess_name_from_resume",
]
